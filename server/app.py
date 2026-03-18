from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
import requests
import mammoth
import psycopg2
from dotenv import load_dotenv
import os
import sys

import pytesseract
from pdf2image import convert_from_path
from PyPDF2 import PdfReader
from PIL import Image


# Fix Python path so we can import model folder
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.append(parent_dir)

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

POPPLER_PATH = r"C:\Users\Sathvik\Downloads\Release-23.11.0-0\poppler-23.11.0\Library\bin"

from model.similarity import get_document
from legalbert import analyze_legal_text

# Load environment variables
load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

app = Flask(__name__)
CORS(app)

# Database connection
db = psycopg2.connect(
    database=os.getenv('DATABASE_NAME'),
    user=os.getenv('DATABASE_USER'),
    password=os.getenv('PASSWORD'),
    host=os.getenv('DATABASE_HOST'),
    port=os.getenv('DATABASE_PORT')
)

# -----------------------------
# Gemini API
# -----------------------------

def gemini_answer(query, context):

    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent"

    headers = {
        "Content-Type": "application/json",
        "x-goog-api-key": GEMINI_API_KEY
    }

    data = {
        "contents": [{
            "parts": [{
                "text": f"""
You are a legal assistant.

Document:
{context}

Question:
{query}

Answer clearly and quote relevant clauses.
"""
            }]
        }]
    }

    response = requests.post(url, headers=headers, json=data)
    result = response.json()

    try:
        return result["candidates"][0]["content"]["parts"][0]["text"]
    except Exception as e:
        
        print("Gemini ERROR:", result)
        print("Exception:", e)
        return "Gemini failed"

# -----------------------------
# Extract DOCX text
# -----------------------------

def extract_text(doc_path):

    doc = Document(doc_path)
    text = []

    for para in doc.paragraphs:
        text.append(para.text)

    return "\n".join(text)

# -----------------------------
# Get Services
# -----------------------------

@app.route('/api/services', methods=["GET"])
def services():

    cur = db.cursor()
    cur.execute("SELECT * FROM services")

    headers = [x[0] for x in cur.description]
    rows = cur.fetchall()

    data = []

    for row in rows:
        data.append(dict(zip(headers, row)))

    cur.close()

    return jsonify(data)

# -----------------------------
# Get Forms
# -----------------------------

@app.route('/api/forms', methods=["GET"])
def get_forms():

    service_id = request.args.get('service_id')

    cur = db.cursor()

    cur.execute("""
    SELECT services.service_id, services.service_name,
           forms.form_id, forms.form_name, forms.form_link
    FROM services
    INNER JOIN forms ON services.service_id = forms.service_id
    WHERE forms.service_id = %s
    """, [service_id])

    headers = [x[0] for x in cur.description]
    rows = cur.fetchall()

    data = []

    for row in rows:
        data.append(dict(zip(headers, row)))

    cur.close()

    return jsonify(data)

# -----------------------------
# Get Form Details
# -----------------------------

@app.route('/api/form-details', methods=["GET"])
def get_form_details():

    form_id = request.args.get('form_id')

    cur = db.cursor()

    cur.execute("SELECT * FROM forms WHERE form_id = %s;", [form_id])
    headers = [x[0] for x in cur.description]
    rows = cur.fetchall()

    data = []

    for row in rows:
        data.append(dict(zip(headers, row)))

    cur.execute("""
    SELECT * FROM ques_categories
    WHERE id IN (
        SELECT DISTINCT(category_id)
        FROM input_ques
        WHERE ques_id IN (
            SELECT form_query_id
            FROM form_queries
            WHERE form_id = %s
        )
    )
    """, [form_id])

    headers = [x[0] for x in cur.description]
    rows = cur.fetchall()

    for row in rows:
        data.append(dict(zip(headers, row)))

    cur.execute("""
    SELECT * FROM input_ques
    WHERE ques_id IN (
        SELECT form_query_id
        FROM form_queries
        WHERE form_id = %s
    )
    """, [form_id])

    headers = [x[0] for x in cur.description]
    rows = cur.fetchall()

    for row in rows:
        data.append(dict(zip(headers, row)))

    cur.close()

    return jsonify(data)

# -----------------------------
# Generate Document
# -----------------------------
@app.route('/api/final-content', methods=["POST"])
def final_content():

    # -----------------------------
    # Receive form data
    # -----------------------------

    form_details = request.json
    print("Form details:", form_details)

    form_details.pop("form_id", None)

    # -----------------------------
    # Load LOCAL template
    # -----------------------------

    template_path = "docs/localfile.docx"
    doc = Document(template_path)

    # -----------------------------
    # Create placeholder mapping
    # -----------------------------

    placeholder_mapping = {}

    for key, value in form_details.items():
        placeholder_mapping[f"#{key}"] = str(value)

    # -----------------------------
    # Function to replace placeholders
    # -----------------------------

    def replace_text(text):

        # replace longer placeholders first (#10 before #1)
        for placeholder in sorted(placeholder_mapping.keys(), key=len, reverse=True):
            text = text.replace(placeholder, placeholder_mapping[placeholder])

        return text

    # -----------------------------
    # Replace in paragraphs
    # -----------------------------

    for paragraph in doc.paragraphs:
        paragraph.text = replace_text(paragraph.text)

    # -----------------------------
    # Replace in tables
    # -----------------------------

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = replace_text(paragraph.text)

    # -----------------------------
    # Save generated document
    # -----------------------------

    output_path = "docs/Output2.docx"

    if os.path.exists(output_path):
        os.remove(output_path)

    doc.save(output_path)

    # -----------------------------
    # Convert DOCX → HTML
    # -----------------------------

    with open(output_path, "rb") as f:
        html = mammoth.convert_to_html(f)

    print("Final document generated")

    return jsonify({
        "content": html.value
    })


# -----------------------------
# Basic Chatbot
# -----------------------------

@app.route('/api/chat', methods=['POST'])
def chat():

    user_input = request.json
    response = get_document(user_input['user_chat'])

    return jsonify({'aiMessage': response})

# -----------------------------
# Chat with Generated Document
# -----------------------------

@app.route('/api/chatdoc', methods=['POST'])
def chatdoc():

    data = request.get_json()

    query = data.get("query")
    doc_name = data.get("doc_name", "Output2.docx")

    doc_path = os.path.join("docs", doc_name)

    if not os.path.exists(doc_path):
        return jsonify({"error": "Document not found"}), 404

    doc_text = extract_text(doc_path)

    # LegalBERT embedding
    embedding = analyze_legal_text(doc_text)

    answer = gemini_answer(query, doc_text)

    return jsonify({
        "answer": answer,
        "embedding_shape": str(embedding.shape)
    })

def extract_text_from_pdf(file_path):

    text = ""

    # 🔹 Try normal extraction first
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() or ""
    except:
        pass

    # 🔹 If text is too small → use OCR
    if len(text.strip()) < 50:
        print("Using OCR for scanned PDF...")

        images = convert_from_path(file_path, poppler_path=POPPLER_PATH)

        for img in images:
            text += pytesseract.image_to_string(img)

    return text
@app.route('/api/pdf-chat', methods=['POST'])
def pdf_chat():

    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    query = request.form.get("query")

    file_path = os.path.join("docs", file.filename)
    file.save(file_path)

    # 🔥 Extract text
    text = extract_text_from_pdf(file_path)

    if len(text.strip()) == 0:
        return jsonify({"error": "Could not extract text"}), 500

    # 🔥 Chunk text (important for large PDFs)
    chunks = [text[i:i+3000] for i in range(0, len(text), 3000)]

    answers = []

    for chunk in chunks[:3]:  # only top 3 chunks
        ans = gemini_answer(query, chunk)
        answers.append(ans)

    final_answer = "\n".join(answers)

    return jsonify({
        "answer": final_answer
    })


# -----------------------------
# Run Server
# -----------------------------

if __name__ == "__main__":
    app.run(debug=True)